from fpdf import FPDF
import os
from docx import Document
from datetime import datetime

# Matriz de Consecuencias (Rule Engine)
MATRIZ = {
    "RULA": {
        "High": {"gravity": "Grave", "action": "Llamado de Atención"},
        "Critical": {"gravity": "Grave/Crítica", "action": "Citación a Descargos"}
    },
    "REBA": {
        "High": {"gravity": "Grave", "action": "Llamado de Atención"},
        "Critical": {"gravity": "Grave/Crítica", "action": "Citación a Descargos"}
    },
    "EPP": {
        "Medium": {"gravity": "Leve", "action": "Nota Verbal / Registro Dashboard"},
        "High": {"gravity": "Grave", "action": "Llamado de Atención"},
        "Critical": {"gravity": "Muy Grave", "action": "Citación a Descargos"}
    }
}

NORMATIVA = "Artículos del Reglamento Interno de Trabajo y Decreto 1072 de 2015 (Seguridad y Salud en el Trabajo)."

def classify_violation(method, risk_level):
    """Clasifica la falta según la matriz de consecuencias."""
    try:
        risk = risk_level.capitalize()
        category = MATRIZ.get(method.upper(), MATRIZ["RULA"])
        return category.get(risk, {"gravity": "Leve", "action": "Notificación Preventiva"})
    except:
        return {"gravity": "Leve", "action": "Notificación Preventiva"}

def create_pdf_report(worker_id, worker_name, worker_full_name, method, risk_level, timestamp, classification, description, evidence_path=""):
    """Genera un reporte oficial en formato PDF con fotografía de evidencia."""
    try:
        print(f"DEBUG: Iniciando generación de PDF para {worker_full_name}")
        pdf = FPDF()
        pdf.add_page()
        
        # Encabezado
        pdf.set_font("Helvetica", "B", 18)
        pdf.set_text_color(44, 62, 80)
        pdf.cell(0, 15, "REPORTE DE INCIDENCIA SST 4.0", 0, 1, "C")
        pdf.ln(5)
        
        # Cuadro de Detalles
        pdf.set_fill_color(240, 240, 240)
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 10, "  DETALLES DEL EVENTO", 0, 1, "L", True)
        
        pdf.set_font("Helvetica", "", 11)
        pdf.set_text_color(0, 0, 0)
        pdf.ln(2)
        pdf.cell(40, 8, "TRABAJADOR:", 0, 0)
        pdf.set_font("Helvetica", "B", 11)
        pdf.cell(0, 8, f"{worker_full_name} (ID: {worker_id})", 0, 1)
        
        pdf.set_font("Helvetica", "", 11)
        pdf.cell(40, 8, "FECHA Y HORA:", 0, 0)
        pdf.cell(0, 8, f"{timestamp}", 0, 1)
        
        pdf.cell(40, 8, "GRAVEDAD:", 0, 0)
        pdf.set_text_color(200, 0, 0)
        pdf.cell(0, 8, f"{classification['gravity'].upper()}", 0, 1)
        
        pdf.set_text_color(0, 0, 0)
        pdf.cell(40, 8, "ACCION:", 0, 0)
        pdf.set_font("Helvetica", "B", 11)
        pdf.cell(0, 8, f"{classification['action']}", 0, 1)
        pdf.ln(5)
        
        # Fotografía de Evidencia
        if evidence_path and os.path.exists(evidence_path):
            pdf.set_font("Helvetica", "B", 12)
            pdf.cell(0, 10, "  EVIDENCIA FOTOGRAFICA", 0, 1, "L", True)
            pdf.ln(5)
            # Insertar imagen (ajustada al ancho de la página)
            pdf.image(evidence_path, x=20, w=170)
            pdf.ln(5)
        
        # Descripción Legal
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 10, "  DESCRIPCION DE LOS HECHOS", 0, 1, "L", True)
        pdf.set_font("Helvetica", "", 11)
        pdf.multi_cell(0, 8, description)
        pdf.ln(10)
        
        # Firmas
        pdf.ln(10)
        pdf.cell(80, 10, "_" * 30, 0, 0, "C")
        pdf.cell(30, 10, "", 0, 0)
        pdf.cell(80, 10, "_" * 30, 0, 1, "C")
        pdf.cell(80, 10, "Firma del Trabajador", 0, 0, "C")
        pdf.cell(30, 10, "", 0, 0)
        pdf.cell(80, 10, "Coordinador SST", 0, 1, "C")
        
        report_dir = "reportes_legales"
        if not os.path.exists(report_dir): os.makedirs(report_dir)
        
        filename = f"{report_dir}/Reporte_{worker_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        pdf.output(filename)
        print(f"DEBUG: PDF con imagen generado: {filename}")
        return filename
    except Exception as e:
        print(f"ERROR: Fallo al generar PDF con imagen: {e}")
        return None

def create_document_draft(worker_id, worker_name, method, risk_level, evidence_path=""):
    """Genera el reporte oficial en PDF con todos los detalles solicitados."""
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    classification = classify_violation(method, risk_level)
    
    # Nombre amigable para el reporte
    worker_full_name = worker_name.replace("_", " ").title()
    
    description = (f"El dia {timestamp}, el sistema de inteligencia artificial detecto un evento de "
                   f"incumplimiento de normas de Seguridad y Salud en el Trabajo. "
                   f"El trabajador fue identificado realizando una postura de riesgo '{risk_level}' "
                   f"utilizando el metodo de analisis {method}. "
                   f"Este documento sirve como notificacion formal y registro de evidencia.")
    
    # Generar PDF (Oficial)
    pdf_path = create_pdf_report(worker_id, worker_name, worker_full_name, method, risk_level, timestamp, classification, description, evidence_path)
    
    if pdf_path:
        # También guardamos el .docx por retrocompatibilidad
        try:
            doc = Document()
            doc.add_heading('NOTIFICACION SST', 0)
            doc.add_paragraph(description)
            if evidence_path: doc.add_picture(evidence_path)
            doc.save(pdf_path.replace(".pdf", ".docx"))
        except: pass
        return pdf_path
    
    return "Error_Reporte"
